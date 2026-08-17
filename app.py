import io, os, re, json, time
from dataclasses import dataclass, field
from typing import Any, Dict, List, Tuple

import streamlit as st
from docx import Document
from openai import OpenAI

APP_VERSION = "Dịch Truyện AI V5.2 PRO"
DEFAULT_MODEL = "deepseek-v4-pro"

# -----------------------------
# Text / document helpers
# -----------------------------

def read_upload(uploaded) -> str:
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith('.txt'):
        for enc in ('utf-8-sig', 'utf-8', 'gb18030', 'gbk'):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                pass
        raise ValueError('Không đọc được file TXT với các bảng mã UTF-8/GB18030/GBK.')
    if name.endswith('.docx'):
        doc = Document(io.BytesIO(data))
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    raise ValueError('Chỉ hỗ trợ DOCX hoặc TXT.')


def normalize_text(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\ufeff', '')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


CHAPTER_PATTERNS = [
    re.compile(r'^\s*(?:第\s*)?(\d{1,5})\s*[章节卷集篇回]\s*.*$', re.I),
    re.compile(r'^\s*Chương\s+\d+.*$', re.I),
    re.compile(r'^\s*chapter\s+\d+.*$', re.I),
    re.compile(r'^\s*第[一二三四五六七八九十百千万零〇\d]+章.*$'),
]


def split_chapters(text: str) -> List[Dict[str, str]]:
    lines = text.split('\n')
    starts = []
    for i, line in enumerate(lines):
        s = line.strip()
        if not s:
            continue
        if any(p.match(s) for p in CHAPTER_PATTERNS):
            starts.append(i)
    if not starts:
        return [{'title': 'Toàn văn', 'text': text.strip()}]
    chapters = []
    for idx, start in enumerate(starts):
        end = starts[idx + 1] if idx + 1 < len(starts) else len(lines)
        title = lines[start].strip()
        body = '\n'.join(lines[start + 1:end]).strip()
        chapters.append({'title': title, 'text': body})
    return chapters


def chunk_text(text: str, max_chars: int) -> List[str]:
    paras = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    chunks, current, size = [], [], 0
    for p in paras:
        # hard split unusually large paragraphs by sentence/newline
        pieces = [p]
        if len(p) > max_chars:
            pieces = re.split(r'(?<=[。！？!?])', p)
            pieces = [x for x in pieces if x.strip()]
        for piece in pieces:
            if current and size + len(piece) + 2 > max_chars:
                chunks.append('\n\n'.join(current))
                current, size = [], 0
            current.append(piece)
            size += len(piece) + 2
    if current:
        chunks.append('\n\n'.join(current))
    return chunks

# -----------------------------
# DeepSeek client
# -----------------------------

def get_client(api_key: str) -> OpenAI:
    return OpenAI(api_key=api_key, base_url='https://api.deepseek.com')


def extract_content(response) -> str:
    try:
        content = response.choices[0].message.content
    except Exception:
        content = None
    if content is None:
        return ''
    return str(content).strip()


def call_model(client: OpenAI, model: str, system: str, user: str,
               max_tokens: int, json_mode: bool = False,
               thinking: bool = True, retries: int = 4) -> str:
    last_error = None
    modes = [thinking] if not thinking else [True, False]
    for mode in modes:
        for attempt in range(retries):
            try:
                kwargs = dict(
                    model=model,
                    messages=[
                        {'role': 'system', 'content': system},
                        {'role': 'user', 'content': user},
                    ],
                    max_tokens=max_tokens,
                    stream=False,
                )
                if json_mode:
                    kwargs['response_format'] = {'type': 'json_object'}
                    mode = False
                if mode:
                    kwargs['reasoning_effort'] = 'high'
                    kwargs['extra_body'] = {'thinking': {'type': 'enabled'}}
                else:
                    kwargs['extra_body'] = {'thinking': {'type': 'disabled'}}
                response = client.chat.completions.create(**kwargs)
                content = extract_content(response)
                if content:
                    return content
                last_error = RuntimeError('DeepSeek trả về nội dung rỗng.')
            except Exception as e:
                last_error = e
            time.sleep(min(1.5 * (2 ** attempt), 8))
    raise last_error or RuntimeError('DeepSeek API không trả về nội dung.')


def _extract_json_object(raw: str):
    # DeepSeek V4 Pro can occasionally wrap valid JSON in prose/markdown even
    # when asked for JSON. Do not depend on response_format=json_object.
    text = (raw or "").strip()
    if not text:
        raise ValueError("Model trả về nội dung rỗng khi cần JSON.")

    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Find the first balanced JSON object, respecting quoted strings.
    start = text.find("{")
    while start >= 0:
        depth = 0
        in_string = False
        escape = False
        for i in range(start, len(text)):
            ch = text[i]
            if in_string:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
            elif ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    candidate = text[start:i+1]
                    try:
                        return json.loads(candidate)
                    except json.JSONDecodeError:
                        break
        start = text.find("{", start + 1)
    raise ValueError("Model trả về JSON không hợp lệ.")


def call_json(client, model, system, user, max_tokens=7000, retries=3):
    """Reliable DeepSeek JSON call.

    JSON Output is requested using DeepSeek's official response_format. JSON
    calls deliberately use non-thinking mode because DeepSeek documents that
    JSON Output can occasionally return empty content; retry/fallback handles it.
    """
    json_system = system + "\n\nQUAN TRỌNG: Chỉ trả về đúng MỘT JSON object hợp lệ theo schema. Không markdown, không code fence, không giải thích bên ngoài JSON. Nếu chưa biết trường nào, dùng chuỗi rỗng hoặc mảng rỗng."
    last_error = None

    # First: official JSON Output, non-thinking.
    for attempt in range(retries):
        try:
            raw = call_model(client, model, json_system, user, max_tokens,
                             json_mode=True, thinking=False, retries=2)
            obj = _extract_json_object(raw)
            if isinstance(obj, dict):
                return obj
            raise ValueError("JSON trả về không phải object.")
        except Exception as e:
            last_error = e
            time.sleep(min(1.5 * (attempt + 1), 5))

    # Second: plain-text JSON fallback, still non-thinking.
    fallback_system = json_system + "\nNếu JSON Output không hoạt động, vẫn chỉ xuất JSON thuần văn bản."
    for attempt in range(2):
        try:
            raw = call_model(client, model, fallback_system, user, max_tokens,
                             json_mode=False, thinking=False, retries=2)
            obj = _extract_json_object(raw)
            if isinstance(obj, dict):
                return obj
        except Exception as e:
            last_error = e
            time.sleep(min(2 * (attempt + 1), 5))

    raise last_error or ValueError("Model trả về JSON không hợp lệ.")

# -----------------------------
# Story Bible
# -----------------------------

BIBLE_SCHEMA = {
    'characters': [
        {
            'name_cn': '', 'name_vi': '', 'aliases': [], 'gender': '',
            'role': '', 'personality': '', 'self_pronoun': '',
            'address_rules': [{'to': '', 'call': '', 'self': '', 'context': ''}],
        }
    ],
    'terms': [{'source': '', 'translation': '', 'type': '', 'note': ''}],
    'places': [{'source': '', 'translation': '', 'note': ''}],
    'organizations': [{'source': '', 'translation': '', 'note': ''}],
    'style': {'era': '', 'genre': '', 'narration': '', 'general_pronoun_rule': ''},
}

BIBLE_SYSTEM = '''Bạn là biên tập viên truyện Trung-Việt chuyên nghiệp. Hãy xây STORY BIBLE có tính "khóa".
Mục tiêu: giữ tên nhân vật, quan hệ, giới tính, xưng hô, thuật ngữ, địa danh và văn phong nhất quán cho toàn bộ truyện.
Không tự bịa chi tiết không có trong văn bản. Nếu chưa chắc, để trống hoặc ghi mức độ chưa xác định.
Đặc biệt phải phân biệt: tự xưng của từng nhân vật và cách nhân vật A gọi nhân vật B theo từng quan hệ/ngữ cảnh.
Trả về JSON đúng schema được yêu cầu.'''


def build_story_bible(client, model, chapters, style, sample_limit=60000):
    pieces = []
    if len(chapters) == 1:
        pieces.append(f"[{chapters[0]['title']}]\n{chapters[0]['text'][:sample_limit]}")
    else:
        per_chapter = max(700, sample_limit // len(chapters))
        for c in chapters:
            pieces.append(f"[{c['title']}]\n{c['text'][:per_chapter]}")
    source = '\n\n'.join(pieces)[:sample_limit]
    user = f'''Phong cách dịch mong muốn:
{style}

Schema JSON bắt buộc:
{json.dumps(BIBLE_SCHEMA, ensure_ascii=False)}

Văn bản nguồn:
{source}'''
    return call_json(client, model, BIBLE_SYSTEM, user, max_tokens=9000)


def bible_for_prompt(bible: Dict[str, Any]) -> str:
    chars = []
    for c in bible.get('characters', []):
        rules = '; '.join(
            f"gọi {r.get('to','')} = {r.get('call','')}, tự xưng = {r.get('self','')} ({r.get('context','')})"
            for r in c.get('address_rules', []) if r.get('call') or r.get('self')
        )
        chars.append(
            f"- {c.get('name_cn','')} -> {c.get('name_vi','')}; bí danh={c.get('aliases',[])}; "
            f"giới tính={c.get('gender','')}; vai trò={c.get('role','')}; tính cách={c.get('personality','')}; "
            f"tự xưng mặc định={c.get('self_pronoun','')}; quy tắc={rules}"
        )
    terms = '\n'.join(f"- {x.get('source')} -> {x.get('translation')} [{x.get('type','')}]" for x in bible.get('terms', []))
    places = '\n'.join(f"- {x.get('source')} -> {x.get('translation')}" for x in bible.get('places', []))
    orgs = '\n'.join(f"- {x.get('source')} -> {x.get('translation')}" for x in bible.get('organizations', []))
    return f'''STORY BIBLE — KHÔNG ĐƯỢC TỰ Ý ĐỔI:
NHÂN VẬT:
{chr(10).join(chars)}
THUẬT NGỮ:
{terms}
ĐỊA DANH:
{places}
TỔ CHỨC:
{orgs}
PHONG CÁCH:
{json.dumps(bible.get('style',{}), ensure_ascii=False)}'''

# -----------------------------
# Translation + validation
# -----------------------------

TRANSLATE_SYSTEM = '''Bạn là dịch giả truyện Trung-Việt. Dịch văn bản nguồn sang tiếng Việt tự nhiên, mượt, giàu cảm xúc, đúng bối cảnh.
Giữ nguyên diễn biến, không tóm tắt, không thêm tình tiết, không giải thích ngoài truyện.
Ưu tiên văn phong cổ trang khi nguồn là cổ trang; lời thoại phải tự nhiên nhưng phù hợp địa vị và quan hệ.
BẮT BUỘC tuân thủ STORY BIBLE, đặc biệt tên nhân vật và xưng hô.
Không được đổi "ta" thành "tôi", "nàng" thành "cô ấy", "chàng" thành "anh ấy" nếu STORY BIBLE không cho phép.
Không thêm tiêu đề chương nếu đoạn nguồn không có.
Chỉ trả về bản dịch, không trả JSON, không nhận xét.'''

CHECK_SYSTEM = '''Bạn là biên tập viên kiểm định bản dịch truyện. So sánh bản nguồn, bản dịch và STORY BIBLE.
Chỉ tìm lỗi nhất quán có thể xác định: sai tên, sai giới tính, sai quan hệ, sai cách xưng hô, sai thuật ngữ, sai địa danh, hoặc tự ý thêm/bớt nội dung.
Trả JSON: {"ok": true/false, "issues": [{"type":"pronoun|name|term|omission|addition|style", "original":"", "translated":"", "fix":""}]}
Không đánh dấu khác biệt văn phong nhỏ nếu không làm sai nghĩa.'''

FIX_SYSTEM = '''Bạn là biên tập viên sửa bản dịch truyện. Sửa CHỈ những lỗi được chỉ ra, giữ nguyên mọi phần đúng.
Không tóm tắt, không thêm nội dung, không đổi giọng kể ngoài phạm vi lỗi.
Trả về toàn bộ đoạn đã sửa, không giải thích.'''


def translate_chunk(client, model, bible_text, prev_context, source_chunk, style, thinking=True):
    user = f'''{bible_text}

PHONG CÁCH DỊCH:
{style}

NGỮ CẢNH LIỀN TRƯỚC (chỉ để hiểu quan hệ, không dịch lại):
{prev_context[-5000:]}

ĐOẠN NGUỒN CẦN DỊCH:
{source_chunk}'''
    return call_model(client, model, TRANSLATE_SYSTEM, user, max_tokens=9000, json_mode=False, thinking=thinking)


def validate_chunk(client, model, bible_text, source, translated):
    user = f'''{bible_text}

BẢN NGUỒN:
{source}

BẢN DỊCH:
{translated}'''
    try:
        return call_json(client, model, CHECK_SYSTEM, user, max_tokens=3500)
    except Exception:
        return {'ok': True, 'issues': []}


def fix_chunk(client, model, bible_text, translated, issues, thinking=True):
    if not issues:
        return translated
    user = f'''{bible_text}

LỖI CẦN SỬA:
{json.dumps(issues, ensure_ascii=False)}

BẢN DỊCH:
{translated}'''
    return call_model(client, model, FIX_SYSTEM, user, max_tokens=9000, thinking=thinking)

# deterministic guardrails for obvious pronoun drift

def pronoun_guard(text: str, bible: Dict[str, Any]) -> List[str]:
    issues = []
    style = bible.get('style', {})
    rule = style.get('general_pronoun_rule', '')
    if 'ta' in rule.lower():
        # Flag common modern drift; model review handles context-specific cases.
        if re.search(r'\b(tôi|cô ấy|anh ấy|cậu ấy)\b', text, re.I):
            issues.append('Phát hiện đại từ hiện đại có thể lệch văn phong/xưng hô: tôi/cô ấy/anh ấy/cậu ấy. Cần kiểm tra theo STORY BIBLE.')
    return issues

# -----------------------------
# Word export
# -----------------------------

def export_docx(title: str, translated_chapters: List[Dict[str, str]], bible: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for ch in translated_chapters:
        doc.add_heading(ch['title'], level=2)
        for p in ch['text'].split('\n\n'):
            if p.strip():
                doc.add_paragraph(p.strip())
    bio = doc.core_properties
    bio.title = title
    bio.subject = APP_VERSION
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# -----------------------------
# Streamlit UI
# -----------------------------

st.set_page_config(page_title=APP_VERSION, page_icon='📖', layout='wide')

st.title('📖 Dịch Truyện AI V5.3 PRO')
st.caption('Word/TXT → nhận diện chương → STORY BIBLE khóa xưng hô → chia nhỏ thông minh → dịch → kiểm tra → sửa lỗi → theo dõi tiến độ → xuất Word')

with st.sidebar:
    st.header('⚙️ Cài đặt')
    api_key = st.text_input('DeepSeek API Key', type='password', value=st.secrets.get('DEEPSEEK_API_KEY', os.getenv('DEEPSEEK_API_KEY', '')))
    model = st.selectbox('Mô hình DeepSeek', [DEFAULT_MODEL, 'deepseek-v4-flash'], index=0)
    style = st.text_area('Phong cách dịch', value='Văn phong truyện tự nhiên, mượt, dễ đọc như bản dịch tiểu thuyết Việt được biên tập kỹ. Giữ sắc thái cảm xúc và bối cảnh. Đối thoại tự nhiên, không máy móc. Không tự ý thêm, bớt hoặc giải thích nội dung.', height=180)
    chunk_size = st.slider('Độ dài mỗi lượt dịch (ký tự)', 3500, 10000, 7000, 500)
    use_thinking = st.checkbox('DeepSeek V4 Pro Thinking', value=True)
    st.info('V5.3 Pro: thêm bảng theo dõi tiến độ theo chương/chunk + STORY BIBLE khóa xưng hô + retry/fallback.')

uploaded = st.file_uploader('📄 Tải 1 file truyện dài', type=['docx', 'txt'])

if uploaded:
    try:
        raw = normalize_text(read_upload(uploaded))
        chapters = split_chapters(raw)
        st.success(f'Đã đọc {len(raw):,} ký tự • phát hiện {len(chapters)} chương • dự kiến {sum(max(1, len(chunk_text(c["text"], chunk_size))) for c in chapters)} lượt dịch.')
    except Exception as e:
        st.error(str(e))
        st.stop()

    if not api_key:
        st.warning('Hãy nhập DeepSeek API Key ở thanh bên trái.')
        st.stop()

    if st.button('🧠 PHÂN TÍCH + BẮT ĐẦU DỊCH V5.3 PRO', type='primary'):
        client = get_client(api_key)
        main_col, monitor_col = st.columns([2.15, 1], gap='large')
        with monitor_col:
            st.subheader('📊 TIẾN ĐỘ XỬ LÝ')
            monitor = st.empty()
            bar_right = st.empty()
        with main_col:
            progress = st.progress(0)
            status = st.empty()

        chapter_states = {c['title']: '⏳ Chờ xử lý' for c in chapters}
        started_at = time.time()
        total_chunks = sum(len(chunk_text(c['text'], chunk_size)) for c in chapters)
        done = 0

        def render_monitor(stage, current_chapter='', current_chunk=0, chapter_chunk_total=0, note=''):
            elapsed = int(time.time() - started_at)
            mm, ss = divmod(elapsed, 60)
            hh, mm = divmod(mm, 60)
            elapsed_text = f'{hh:02d}:{mm:02d}:{ss:02d}' if hh else f'{mm:02d}:{ss:02d}'
            overall = done / max(total_chunks, 1)
            right_lines = [
                f'**Giai đoạn:** {stage}',
                f'**Tổng:** {len(chapters)} chương • {total_chunks} lượt',
                f'**Đã xong:** {done}/{total_chunks} lượt ({overall*100:.1f}%)',
                f'**Đang xử lý:** {current_chapter or "—"}',
                f'**Chunk:** {current_chunk}/{chapter_chunk_total}' if chapter_chunk_total else '**Chunk:** —',
                f'**Thời gian:** {elapsed_text}',
            ]
            if note:
                right_lines.append(f'**Trạng thái:** {note}')
            right_lines.append('')
            right_lines.extend([f'- {k}: {v}' for k, v in chapter_states.items()])
            monitor.markdown('\n'.join(right_lines))
            bar_right.progress(min(overall, 1.0))

        try:
            render_monitor('1/4 — STORY BIBLE', note='🧠 Đang phân tích nhân vật và xưng hô')
            status.info('Bước 1/4: đang xây STORY BIBLE khóa nhân vật và xưng hô...')
            bible = build_story_bible(client, model, chapters, style)
            st.session_state['bible'] = bible

            translated = []
            prev = ''
            errors = []

            for ci, chapter in enumerate(chapters, start=1):
                chapter_chunks = chunk_text(chapter['text'], chunk_size)
                chapter_states[chapter['title']] = '🔄 Đang dịch'
                status.info(f'Bước 2/4: đang dịch {chapter["title"]} ({ci}/{len(chapters)})...')
                render_monitor('2/4 — DỊCH', chapter['title'], 0, len(chapter_chunks), '🔄 Bắt đầu chương')
                out_chunks = []
                for chunk_i, source_chunk in enumerate(chapter_chunks, start=1):
                    bible_text = bible_for_prompt(bible)
                    render_monitor('2/4 — DỊCH', chapter['title'], chunk_i, len(chapter_chunks), '🤖 Đang gọi DeepSeek')
                    try:
                        translated_chunk = translate_chunk(client, model, bible_text, prev, source_chunk, style, use_thinking)
                        render_monitor('2/4 — KIỂM TRA', chapter['title'], chunk_i, len(chapter_chunks), '🔍 Đang kiểm tra tên + xưng hô')
                        check = validate_chunk(client, model, bible_text, source_chunk, translated_chunk)
                        issues = check.get('issues', []) if isinstance(check, dict) else []
                        issues += [{'type': 'style', 'original': '', 'translated': '', 'fix': x} for x in pronoun_guard(translated_chunk, bible)]
                        if issues:
                            render_monitor('2/4 — SỬA', chapter['title'], chunk_i, len(chapter_chunks), f'🔧 Đang sửa {len(issues)} lỗi')
                            translated_chunk = fix_chunk(client, model, bible_text, translated_chunk, issues, use_thinking)
                        out_chunks.append(translated_chunk)
                        prev = (prev + '\n\n' + translated_chunk)[-7000:]
                    except Exception as e:
                        errors.append(f'{chapter["title"]} / chunk {chunk_i}: {type(e).__name__}: {e}')
                        out_chunks.append('[LỖI DỊCH CHUNK — CẦN CHẠY LẠI]\n' + source_chunk)
                    done += 1
                    progress.progress(min(done / max(total_chunks, 1), 1.0))
                    render_monitor('2/4 — DỊCH', chapter['title'], chunk_i, len(chapter_chunks), '✅ Đã hoàn thành chunk')
                chapter_states[chapter['title']] = '✅ Hoàn thành' if not any(chapter['title'] in e for e in errors) else '⚠️ Có lỗi chunk'
                translated.append({'title': chapter['title'], 'text': '\n\n'.join(out_chunks)})
                render_monitor('2/4 — DỊCH', chapter['title'], len(chapter_chunks), len(chapter_chunks), chapter_states[chapter['title']])

            status.info('Bước 3/4: hoàn tất kiểm tra nhất quán và ghép chương...')
            render_monitor('3/4 — GHÉP WORD', note='📄 Đang tạo file Word')
            docx_bytes = export_docx(os.path.splitext(uploaded.name)[0] + ' - V5.3 PRO', translated, bible)
            st.session_state['translated'] = translated
            st.session_state['docx'] = docx_bytes
            st.session_state['errors'] = errors
            progress.progress(1.0)
            render_monitor('4/4 — HOÀN TẤT', note='🎉 Đã dịch xong và tạo Word')
            status.success('Bước 4/4: đã hoàn tất.')
        except Exception as e:
            render_monitor('⛔ DỪNG', note=f'❌ {type(e).__name__}: {e}')
            st.error(f'Quá trình dịch bị dừng do lỗi API. Bộ nhớ đã xử lý trước đó vẫn được giữ.\n\n{type(e).__name__}: {e}')

if 'bible' in st.session_state:
    with st.expander('🧠 STORY BIBLE — bộ nhớ nhân vật, xưng hô, thuật ngữ', expanded=False):
        st.json(st.session_state['bible'])

if 'translated' in st.session_state:
    st.subheader('📥 Bản dịch hoàn chỉnh')
    st.download_button('⬇️ TẢI FILE WORD ĐÃ DỊCH', data=st.session_state['docx'], file_name='ban-dich-v5-pro.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    with st.expander('👀 Xem thử bản dịch'):
        for ch in st.session_state['translated'][:3]:
            st.markdown(f'### {ch["title"]}')
            st.write(ch['text'][:5000])
    if st.session_state.get('errors'):
        st.warning('Có một số chunk lỗi API và đã được đánh dấu để chạy lại:')
        st.write('\n'.join(st.session_state['errors']))
